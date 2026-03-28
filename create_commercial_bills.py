"""
Create 3 quarterly commercial electricity bill .docx fixtures for the bill intelligence pipeline.
Austin, TX commercial building — TexStar Energy Services
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
PROVIDER = "TexStar Energy Services"
CUSTOMER = "Greenfield Commercial LLC"
ADDRESS = "1200 Industrial Blvd, Austin, TX 78741"
ACCOUNT = "ACC-TX-00829341"
METER = "MTR-TX-551902"
CURRENCY = "USD"

ENERGY_RATE = 0.1100
DELIVERY_RATE = 0.0250
FIXED = 35.00
TAX_RATE = 0.0825

OUTDIR = Path(__file__).parent / "tests" / "fixtures"


# ---------------------------------------------------------------------------
# Charge helpers
# ---------------------------------------------------------------------------
def compute(kwh: int):
    energy = kwh * ENERGY_RATE
    delivery = kwh * DELIVERY_RATE
    subtotal = energy + delivery + FIXED
    tax = round(subtotal * TAX_RATE, 2)
    total = subtotal + tax
    return dict(energy=energy, delivery=delivery, fixed=FIXED,
                subtotal=subtotal, tax=tax, total=total)


# ---------------------------------------------------------------------------
# Quarter definitions
# ---------------------------------------------------------------------------
Q1_KWH = 27000
Q2_KWH = 41000
Q3_KWH = 27000

Q1 = compute(Q1_KWH)
Q2 = compute(Q2_KWH)
Q3 = compute(Q3_KWH)

# Usage history rows: made-up plausible prior quarters + actual
# Prior quarters (roughly 8k-14k/month → 24k-42k/quarter)
PRIOR_QUARTERS = [
    ("2024-Q1", 26200, 3871.93),
    ("2024-Q2", 39800, 5848.04),
    ("2024-Q3", 26500, 3916.30),
    ("2024-Q4", 25800, 3813.21),
    ("2025-Q1-prev", 27100, 3997.34),
]

BILLS = [
    dict(
        name="demo_bill_04_commercial_q1.docx",
        bill_num="TXS-2025-Q1-00441",
        start="2025-01-01",
        end="2025-04-30",
        days=120,
        issue="2025-05-05",
        kwh=Q1_KWH,
        prev_paid=0.00,
        charges=Q1,
        history=[
            ("2023-Q4", 25100, 3712.47),
            ("2024-Q1", 26200, 3871.93),
            ("2024-Q2", 39800, 5848.04),
            ("2024-Q3", 26500, 3916.30),
            ("2024-Q4", 25800, 3813.21),
            ("2025-Q1", Q1_KWH, Q1["total"]),
        ],
    ),
    dict(
        name="demo_bill_05_commercial_q2.docx",
        bill_num="TXS-2025-Q2-00442",
        start="2025-05-01",
        end="2025-08-31",
        days=123,
        issue="2025-09-03",
        kwh=Q2_KWH,
        prev_paid=Q1["total"],
        charges=Q2,
        history=[
            ("2024-Q1", 26200, 3871.93),
            ("2024-Q2", 39800, 5848.04),
            ("2024-Q3", 26500, 3916.30),
            ("2024-Q4", 25800, 3813.21),
            ("2025-Q1", Q1_KWH, Q1["total"]),
            ("2025-Q2", Q2_KWH, Q2["total"]),
        ],
    ),
    dict(
        name="demo_bill_06_commercial_q3.docx",
        bill_num="TXS-2025-Q3-00443",
        start="2025-09-01",
        end="2025-12-31",
        days=122,
        issue="2026-01-05",
        kwh=Q3_KWH,
        prev_paid=Q2["total"],
        charges=Q3,
        history=[
            ("2024-Q3", 26500, 3916.30),
            ("2024-Q4", 25800, 3813.21),
            ("2025-Q1", Q1_KWH, Q1["total"]),
            ("2025-Q2", Q2_KWH, Q2["total"]),
            ("2025-Q3", Q3_KWH, Q3["total"]),
        ],
    ),
]


# ---------------------------------------------------------------------------
# Docx builder
# ---------------------------------------------------------------------------
def add_table(doc: Document, rows: list[list[str]], style: str = "Table Grid"):
    if not rows:
        return
    ncols = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = style
    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            t.rows[r_idx].cells[c_idx].text = str(text)


def build_bill(bill: dict, outdir: Path):
    doc = Document()

    # Remove default empty paragraph at top if present
    # Header table
    header_rows = [
        [f"{PROVIDER}\nElectric Service Statement", "DEMO SAMPLE\nFICTIONAL BILL\nNOT FOR PAYMENT"],
    ]
    add_table(doc, header_rows)

    doc.add_paragraph("")
    doc.add_paragraph("Account Summary")

    c = bill["charges"]
    kwh = bill["kwh"]

    # Main info table
    info_rows = [
        ["Provider Name", PROVIDER, "Bill Number", bill["bill_num"]],
        ["Account Number", ACCOUNT, "Customer Name", CUSTOMER],
        ["Service Address", ADDRESS, "Meter Number", METER],
        ["Billing Period Start", bill["start"], "Billing Period End", bill["end"]],
        ["Billing Days", str(bill["days"]), "Issue Date", bill["issue"]],
        ["Electricity Import (kWh)", str(kwh), "Electricity Export (kWh)", "0"],
        ["Import Rate per kWh", "0.1100", "Previous Amount Paid", f"{bill['prev_paid']:.2f}"],
        ["Currency", "USD", "Total Amount Due", f"{c['total']:.2f}"],
    ]
    add_table(doc, info_rows)

    doc.add_paragraph("")
    doc.add_paragraph("Charges and Credits")

    # Charges table
    charges_rows = [
        ["Line Item", "Quantity / Rate", "Amount (USD)"],
        ["Energy Charge", f"{kwh} kWh x 0.1100", f"{c['energy']:.2f}"],
        ["Delivery Charge", f"{kwh} kWh x 0.0250", f"{c['delivery']:.2f}"],
        ["Fixed Service Charge", "Quarterly fee", "35.00"],
        ["Taxes and Fees", "TX utility tax (8.25%)", f"{c['tax']:.2f}"],
        ["Total Amount Due", "Total Amount Due", f"{c['total']:.2f}"],
    ]
    add_table(doc, charges_rows)

    doc.add_paragraph("")
    doc.add_paragraph("Recent Usage History")

    # Usage history table
    history_rows = [["Period", "Import kWh", "Bill Total (USD)"]]
    for period, h_kwh, h_total in bill["history"]:
        history_rows.append([period, str(h_kwh), f"{h_total:.2f}"])
    add_table(doc, history_rows)

    # Disclaimer paragraph
    doc.add_paragraph("This document is synthetic and intended only for parser and demo testing.")
    doc.add_paragraph("")

    out_path = outdir / bill["name"]
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    OUTDIR.mkdir(parents=True, exist_ok=True)

    print("Computing charges:")
    print(f"  Q1: {Q1_KWH} kWh → energy={Q1['energy']:.2f}, delivery={Q1['delivery']:.2f}, "
          f"fixed={Q1['fixed']:.2f}, tax={Q1['tax']:.2f}, TOTAL={Q1['total']:.2f}")
    print(f"  Q2: {Q2_KWH} kWh → energy={Q2['energy']:.2f}, delivery={Q2['delivery']:.2f}, "
          f"fixed={Q2['fixed']:.2f}, tax={Q2['tax']:.2f}, TOTAL={Q2['total']:.2f}")
    print(f"  Q3: {Q3_KWH} kWh → energy={Q3['energy']:.2f}, delivery={Q3['delivery']:.2f}, "
          f"fixed={Q3['fixed']:.2f}, tax={Q3['tax']:.2f}, TOTAL={Q3['total']:.2f}")
    print(f"  Annual kWh: {Q1_KWH + Q2_KWH + Q3_KWH}")
    print(f"  Annual cost: {Q1['total'] + Q2['total'] + Q3['total']:.2f}")
    print()

    for bill in BILLS:
        build_bill(bill, OUTDIR)

    print("\nDone.")
